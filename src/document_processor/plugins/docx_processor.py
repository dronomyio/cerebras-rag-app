#!/usr/bin/env python3
"""
DOCX Document Processor Plugin
------------------------------
Plugin for processing Microsoft Word documents.
"""

import os
import re
import logging
from typing import Dict, List, Any, Optional
from pathlib import Path

# Import base processor
try:
    from .base_processor import BaseDocumentProcessor
except ImportError:
    # Fallback for when module structure isn't recognized
    import sys
    import os
    # Absolute import as a last resort
    sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    from plugins.base_processor import BaseDocumentProcessor

logger = logging.getLogger(__name__)

class DocxProcessor(BaseDocumentProcessor):
    """
    Processor for Microsoft Word (DOCX) documents.
    """
    
    def __init__(self, config: Dict[str, Any]):
        """
        Initialize the DOCX processor with configuration.
        
        Args:
            config: Configuration dictionary from the main config file
        """
        super().__init__(config)
        
    def can_process(self, file_path: str) -> bool:
        """
        Check if this processor can handle the given file.
        
        Args:
            file_path: Path to the file to check
            
        Returns:
            True if this processor can handle the file, False otherwise
        """
        if not self.is_enabled():
            return False
            
        file_path = Path(file_path)
        return file_path.suffix.lower() in ['.docx', '.doc']
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process the DOCX document and return chunks.
        
        Args:
            file_path: Path to the DOCX file to process
            
        Returns:
            List of chunks, where each chunk is a dictionary with at least
            'content' and 'metadata' keys
        """
        try:
            # Import docx here to avoid dependency issues if not installed
            import docx
            
            file_path = Path(file_path)
            chunks = []
            
            # Extract text from DOCX
            doc = docx.Document(file_path)
            
            # Get document metadata
            metadata = {
                "source": file_path.name,
                "file_type": "docx",
                "processor": "docx_processor"
            }
            
            # Extract document properties if available
            try:
                core_properties = doc.core_properties
                metadata["title"] = core_properties.title or ""
                metadata["author"] = core_properties.author or ""
                metadata["subject"] = core_properties.subject or ""
                metadata["keywords"] = core_properties.keywords or ""
            except:
                logger.warning(f"Could not extract core properties from {file_path}")
            
            # Process with chunking
            current_chunk = {"content": "", "metadata": metadata.copy()}
            current_chunk_size = 0
            chunk_size = self.get_chunk_size()
            chunk_overlap = self.get_chunk_overlap()
            current_heading = None
            
            # Process headings and paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                # Skip empty paragraphs
                if not text:
                    continue
                
                # Check if this is a heading
                if paragraph.style.name.startswith('Heading'):
                    # Start a new chunk for headings
                    if current_chunk["content"]:
                        chunks.append(current_chunk)
                    
                    current_heading = text
                    heading_level = int(paragraph.style.name.replace('Heading ', '')) if paragraph.style.name != 'Heading' else 1
                    
                    current_chunk = {
                        "content": text,
                        "metadata": metadata.copy()
                    }
                    current_chunk["metadata"]["heading"] = text
                    current_chunk["metadata"]["heading_level"] = heading_level
                    current_chunk["metadata"]["element_type"] = "heading"
                    current_chunk_size = len(text)
                    
                else:
                    # If adding this paragraph would exceed chunk size, save current chunk and start a new one
                    if current_chunk_size + len(text) > chunk_size and current_chunk["content"]:
                        chunks.append(current_chunk)
                        
                        # Start new chunk with overlap
                        overlap_text = current_chunk["content"][-chunk_overlap:] if chunk_overlap > 0 else ""
                        current_chunk = {
                            "content": overlap_text,
                            "metadata": metadata.copy()
                        }
                        if current_heading:
                            current_chunk["metadata"]["heading"] = current_heading
                        current_chunk_size = len(overlap_text)
                    
                    # Add paragraph to current chunk
                    if current_chunk["content"]:
                        current_chunk["content"] += "\n\n"
                    current_chunk["content"] += text
                    current_chunk_size = len(current_chunk["content"])
            
            # Add the last chunk if not empty
            if current_chunk["content"]:
                chunks.append(current_chunk)
            
            # Process tables if configured
            if self.get_setting("extract_tables", True):
                for i, table in enumerate(doc.tables):
                    table_text = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        table_text.append(" | ".join(row_text))
                    
                    table_content = "\n".join(table_text)
                    if table_content.strip():
                        table_metadata = metadata.copy()
                        table_metadata["element_type"] = "table"
                        table_metadata["table_index"] = i
                        if current_heading:
                            table_metadata["heading"] = current_heading
                        
                        chunks.append({
                            "content": table_content,
                            "metadata": table_metadata
                        })
            
            return chunks
            
        except ImportError:
            logger.error("python-docx is not installed. Cannot process DOCX.")
            return [{
                "content": "Error: python-docx is not installed. Cannot process DOCX.",
                "metadata": {
                    "source": Path(file_path).name,
                    "file_type": "docx",
                    "error": "Missing dependency: python-docx"
                }
            }]
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            # Return a single chunk with error information
            return [{
                "content": f"Error processing DOCX: {str(e)}",
                "metadata": {
                    "source": Path(file_path).name,
                    "file_type": "docx",
                    "error": str(e)
                }
            }]
